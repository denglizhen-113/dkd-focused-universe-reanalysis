#!/usr/bin/env python3
"""Build the complete M19 Scientific Reports submission package from frozen tables."""

from __future__ import annotations

import csv
import hashlib
import json
import shutil
import subprocess
import sys
import textwrap
import zipfile
from pathlib import Path

import numpy as np
import pandas as pd


ROOT = Path(__file__).resolve().parents[2]
TABLES = ROOT / "tables" / "stage21_m19_scientific_reports_revision"
DOCS = ROOT / "docs" / "stage21_m19_scientific_reports_revision"
MANUSCRIPT = ROOT / "manuscript_ready" / "stage21_m19_scientific_reports_revision"
PACKAGE = ROOT / "submission_package" / "stage21_m19_scientific_reports_revision"
SUPP = PACKAGE / "supplementary"
MAIN_FIG = PACKAGE / "main_figures"
READY = ROOT / "submission_ready_scientific_reports_m19"
PANDOC = Path(r"C:\ProgramData\anaconda3\Library\bin\pandoc.exe")
PDF_EXPORTER = ROOT / "scripts" / "stage21_m17_assembly" / "export_m17_pdf.ps1"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def markdown_table(frame: pd.DataFrame, columns: list[str], labels: list[str]) -> str:
    data = frame[columns].copy().fillna("")
    data.columns = labels
    lines = ["| " + " | ".join(labels) + " |", "| " + " | ".join(["---"] * len(labels)) + " |"]
    for row in data.astype(str).itertuples(index=False, name=None):
        lines.append("| " + " | ".join(value.replace("|", "/") for value in row) + " |")
    return "\n".join(lines)


def build_figures() -> None:
    import matplotlib.pyplot as plt
    from matplotlib.colors import LinearSegmentedColormap
    from matplotlib.lines import Line2D
    from matplotlib.patches import FancyBboxPatch, Patch

    # Double-column Scientific Reports layout (183 mm wide), with conservative
    # journal-scale typography and embedded TrueType fonts in vector output.
    plt.rcParams.update({
        "font.family": "Arial",
        "font.sans-serif": ["Arial", "Liberation Sans", "DejaVu Sans"],
        "font.size": 7.2,
        "axes.labelsize": 7.5,
        "axes.titlesize": 8.2,
        "axes.titleweight": "bold",
        "xtick.labelsize": 6.8,
        "ytick.labelsize": 6.8,
        "legend.fontsize": 6.5,
        "axes.linewidth": 0.7,
        "axes.spines.top": False,
        "axes.spines.right": False,
        "pdf.fonttype": 42,
        "ps.fonttype": 42,
        "savefig.facecolor": "white",
        "figure.facecolor": "white",
    })
    MAIN_FIG.mkdir(parents=True, exist_ok=True)
    SUPP.mkdir(parents=True, exist_ok=True)

    navy = "#0072B2"       # Okabe-Ito blue
    sky = "#56B4E9"        # Okabe-Ito sky blue
    orange = "#E69F00"     # Okabe-Ito orange
    green = "#009E73"      # Okabe-Ito bluish green
    vermilion = "#D55E00"  # Okabe-Ito vermilion
    purple = "#CC79A7"     # Okabe-Ito reddish purple
    charcoal = "#2F3337"
    midgrey = "#7A7F85"
    lightgrey = "#E6E8EA"
    pale = "#F5F7F8"

    def panel_label(ax, label: str, x: float = -0.08, y: float = 1.03) -> None:
        ax.text(x, y, label, transform=ax.transAxes, fontsize=8, fontweight="bold",
                va="bottom", ha="left", color="black", clip_on=False)

    def save_main(fig, stem: str) -> None:
        fig.savefig(MAIN_FIG / f"{stem}.png", dpi=450, bbox_inches="tight", pad_inches=0.04)
        fig.savefig(MAIN_FIG / f"{stem}.pdf", bbox_inches="tight", pad_inches=0.04,
                    metadata={"Creator": "Matplotlib; generated from archived numerical tables"})
        plt.close(fig)

    # Figure 1: compact PRISMA-style dataset-level flow with exclusion branches.
    fig, ax = plt.subplots(figsize=(7.20, 5.25))
    ax.axis("off")
    main_x, side_x = 0.35, 0.77
    box_w, main_h, side_w = 0.48, 0.115, 0.34
    stages = [
        (0.89, "Records identified", "GEO, n = 263\nArrayExpress, n = 54\nPubMed-derived accessions, n = 57\nTotal, n = 374"),
        (0.69, "Records screened", "After deduplication, n = 322"),
        (0.49, "Full records assessed", "Accession and sample metadata reviewed, n = 53"),
        (0.25, "Datasets included", "11 GEO Series representing\n9 independent source studies"),
        (0.075, "Primary synthesis", "3 independent glomerular sources"),
    ]
    for y0, heading, body in stages:
        h = 0.145 if y0 == 0.89 else main_h
        patch = FancyBboxPatch((main_x - box_w / 2, y0 - h / 2), box_w, h,
                               boxstyle="round,pad=0.012,rounding_size=0.01",
                               transform=ax.transAxes, facecolor=pale, edgecolor=navy, linewidth=1.0)
        ax.add_patch(patch)
        ax.text(main_x, y0 + (0.022 if y0 == 0.89 else 0.018), heading,
                transform=ax.transAxes, ha="center", va="center", fontsize=7.4,
                fontweight="bold", color=charcoal)
        ax.text(main_x, y0 - (0.035 if y0 == 0.89 else 0.022), body,
                transform=ax.transAxes, ha="center", va="center", fontsize=6.7,
                color=charcoal, linespacing=1.2)
    exclusions = [
        (0.79, "Records removed before screening", "Duplicates, n = 52"),
        (0.49, "Excluded at title/summary", "n = 269"),
        (0.31, "Excluded after full review", "n = 42\nIneligible contrast/specimen/assay,\nexperimental model, or duplicate source"),
    ]
    for y0, heading, body in exclusions:
        h = 0.105 if y0 != 0.31 else 0.145
        patch = FancyBboxPatch((side_x - side_w / 2, y0 - h / 2), side_w, h,
                               boxstyle="round,pad=0.01,rounding_size=0.01",
                               transform=ax.transAxes, facecolor="white", edgecolor=midgrey, linewidth=0.8)
        ax.add_patch(patch)
        ax.text(side_x, y0 + 0.018, heading, transform=ax.transAxes, ha="center", va="center",
                fontsize=6.8, fontweight="bold", color=charcoal)
        ax.text(side_x, y0 - 0.023, body, transform=ax.transAxes, ha="center", va="center",
                fontsize=6.2, color=charcoal, linespacing=1.15)
    for y1, y2 in ((0.81, 0.755), (0.625, 0.555), (0.425, 0.325), (0.19, 0.135)):
        ax.annotate("", xy=(main_x, y2), xytext=(main_x, y1), xycoords="axes fraction",
                    arrowprops=dict(arrowstyle="-|>", mutation_scale=8, lw=0.9, color=navy))
    for y0, start_y in ((0.79, 0.82), (0.49, 0.49), (0.31, 0.31)):
        ax.annotate("", xy=(side_x - side_w / 2, y0), xytext=(main_x + box_w / 2, start_y),
                    xycoords="axes fraction", arrowprops=dict(arrowstyle="-|>", mutation_scale=8,
                                                              lw=0.8, color=midgrey))
    for y0, label in ((0.89, "Identification"), (0.69, "Screening"),
                      (0.49, "Eligibility"), (0.25, "Included")):
        ax.text(0.035, y0, label, transform=ax.transAxes, rotation=90, ha="center", va="center",
                fontsize=6.4, fontweight="bold", color=midgrey)
    ax.text(0.01, 0.99, "Dataset-level flow", transform=ax.transAxes, ha="left", va="top",
            fontsize=8.2, fontweight="bold")
    ax.text(0.99, 0.99, "Search completed 21 August 2026", transform=ax.transAxes,
            ha="right", va="top", fontsize=6.4, color=midgrey)
    save_main(fig, "Figure_1")

    # Figure 2: study architecture. Panel a shows group sizes; panel b makes
    # source-study overlap and compartment separation visually explicit.
    cohorts = pd.read_csv(TABLES / "cohort_characteristics.csv")
    compartment_order = ["glomerular", "tubulointerstitial", "whole/cortical kidney", "interstitium only"]
    compartment_label = {"glomerular": "Glomerular", "tubulointerstitial": "Tubulointerstitial",
                         "whole/cortical kidney": "Whole/cortical", "interstitium only": "Interstitium"}
    compartment_color = {"glomerular": navy, "tubulointerstitial": orange,
                         "whole/cortical kidney": green, "interstitium only": purple}
    role_marker = {"primary glomerular meta-analysis": "o",
                   "compartment-specific contextual analysis": "s",
                   "small-sample glomerular sensitivity": "^"}
    source_order = list(dict.fromkeys(cohorts["source_study"].tolist()))
    cohorts = cohorts.copy()
    cohorts["_source_rank"] = cohorts["source_study"].map({s: i for i, s in enumerate(source_order)})
    cohorts["_compartment_rank"] = cohorts["compartment"].map({s: i for i, s in enumerate(compartment_order)})
    cohorts = cohorts.sort_values(["_source_rank", "_compartment_rank"]).reset_index(drop=True)
    fig = plt.figure(figsize=(6.92, 6.05), constrained_layout=True)
    gs = fig.add_gridspec(1, 2, width_ratios=[1.45, 1.0], wspace=0.16)
    ax = fig.add_subplot(gs[0, 0])
    axm = fig.add_subplot(gs[0, 1])
    y = np.arange(len(cohorts))[::-1]
    ax.barh(y, -cohorts["n_control"], height=0.64, color="#B8BDC2", edgecolor="white", linewidth=0.4)
    ax.barh(y, cohorts["n_case"], height=0.64,
            color=[compartment_color[x] for x in cohorts["compartment"]], edgecolor="white", linewidth=0.4)
    ax.axvline(0, color=charcoal, lw=0.65)
    ax.set_yticks(y, cohorts["cohort"])
    ax.set_xlim(-22, 44)
    ax.set_xticks([-20, -10, 0, 10, 20, 30, 40], [20, 10, 0, 10, 20, 30, 40])
    ax.set_xlabel("Number of samples")
    ax.set_title("Cohort group sizes", loc="left", pad=5)
    ax.grid(axis="x", color=lightgrey, linewidth=0.45, zorder=0)
    ax.set_axisbelow(True)
    for yi, row in zip(y, cohorts.itertuples()):
        ax.text(-row.n_control - 0.5, yi, f"{row.n_control}", ha="right", va="center", fontsize=6.0, color=charcoal)
        ax.text(row.n_case + 0.5, yi, f"{row.n_case}", ha="left", va="center", fontsize=6.0, color=charcoal)
    ax.text(0.29, 1.005, "Control", transform=ax.transAxes, ha="center", va="bottom", color=midgrey, fontsize=6.4)
    ax.text(0.70, 1.005, "DKD", transform=ax.transAxes, ha="center", va="bottom", color=charcoal, fontsize=6.4)
    panel_label(ax, "a", x=-0.18)

    # One row per independent source study. Two cells on one row reveal Series
    # that derive from the same source and therefore cannot be counted twice.
    source_y = {s: len(source_order) - 1 - i for i, s in enumerate(source_order)}
    comp_x = {c: i for i, c in enumerate(compartment_order)}
    for row in cohorts.itertuples():
        x0, y0 = comp_x[row.compartment], source_y[row.source_study]
        axm.scatter(x0, y0, s=28 + 1.6 * (row.n_case + row.n_control),
                    marker=role_marker[row.analysis_tier], facecolor=compartment_color[row.compartment],
                    edgecolor="white", linewidth=0.7, zorder=3)
        short = row.cohort.replace("_advanced", "").replace("_donor_averaged", "").replace("_interstitium", "")
        axm.text(x0, y0 - 0.29, short.replace("GSE", ""), ha="center", va="top", fontsize=5.3, color=charcoal)
    axm.set_xlim(-0.55, 3.55); axm.set_ylim(-0.7, len(source_order) - 0.3)
    axm.set_xticks(range(4), ["Glom.", "Tubulo-\ninterstitial", "Whole/\ncortical", "Interstitium"])
    axm.set_yticks(range(len(source_order)), list(reversed(source_order)))
    axm.set_title("Sources and compartments", loc="left", pad=5)
    axm.grid(color=lightgrey, linewidth=0.45)
    axm.tick_params(axis="both", length=0)
    for spine in axm.spines.values():
        spine.set_visible(False)
    role_handles = [
        Line2D([0], [0], marker="o", color="none", markerfacecolor=midgrey, markeredgecolor="white", markersize=6, label="Primary synthesis"),
        Line2D([0], [0], marker="s", color="none", markerfacecolor=midgrey, markeredgecolor="white", markersize=6, label="Contextual"),
        Line2D([0], [0], marker="^", color="none", markerfacecolor=midgrey, markeredgecolor="white", markersize=6, label="Sensitivity"),
    ]
    axm.legend(handles=role_handles, loc="lower center", bbox_to_anchor=(0.5, -0.20),
               ncol=1, frameon=False, handletextpad=0.4, borderaxespad=0)
    compartment_handles = [Patch(facecolor=compartment_color[c], edgecolor="none", label=compartment_label[c])
                           for c in compartment_order]
    ax.legend(handles=compartment_handles, loc="lower center", bbox_to_anchor=(0.50, -0.17),
              ncol=2, frameon=False, handlelength=1.0, columnspacing=0.8, borderaxespad=0)
    panel_label(axm, "b", x=-0.16)
    save_main(fig, "Figure_2")

    # Figure 3: a conventional selected-gene forest plus an honest all-gene
    # multiplicity landscape. The selection is explicitly based on raw P only.
    meta = pd.read_csv(TABLES / "primary_glomerular_gene_meta.csv")
    complete = meta.loc[meta["k"].eq(3) & meta["p_value_modified_hk"].notna()].copy()
    display = complete.nsmallest(12, "p_value_modified_hk").sort_values("pooled_effect").reset_index(drop=True)
    fig = plt.figure(figsize=(7.20, 5.15), constrained_layout=True)
    gs = fig.add_gridspec(1, 2, width_ratios=[1.38, 1.0], wspace=0.12)
    ax = fig.add_subplot(gs[0, 0])
    axb = fig.add_subplot(gs[0, 1])
    yy = np.arange(len(display))
    x = display["pooled_effect"].to_numpy()
    lo = display["ci_95_low_modified_hk"].to_numpy()
    hi = display["ci_95_high_modified_hk"].to_numpy()
    for yi in range(len(display)):
        if yi % 2 == 0:
            ax.axhspan(yi - 0.5, yi + 0.5, color=pale, zorder=0)
    point_colors = [vermilion if value > 0 else navy for value in x]
    ax.errorbar(x, yy, xerr=[x - lo, hi - x], fmt="none", ecolor="#788E9F",
                elinewidth=1.0, capsize=2.2, capthick=0.8, zorder=2)
    ax.scatter(x, yy, s=28, c=point_colors, edgecolor="white", linewidth=0.5, zorder=3)
    ax.axvline(0, color=charcoal, lw=0.75)
    ax.set_yticks(yy, display["gene_symbol"])
    ax.set_ylim(-0.7, len(display) - 0.3)
    xmin = min(-2.4, np.nanmin(lo) - 0.15); xmax = max(3.2, np.nanmax(hi) + 0.15)
    ax.set_xlim(xmin, xmax)
    ax.set_xlabel("Pooled Hedges' g (95% modified HK CI)")
    ax.set_title("Lowest unadjusted P values", loc="left", pad=5)
    ax.grid(axis="x", color=lightgrey, linewidth=0.45, zorder=0)
    panel_label(ax, "a", x=-0.18)
    ax.text(0.99, 0.01, "Displayed genes are descriptive only", transform=ax.transAxes,
            ha="right", va="bottom", fontsize=5.8, color=midgrey)

    px = complete["pooled_effect"].to_numpy()
    py = -np.log10(np.clip(complete["p_value_modified_hk"].to_numpy(), np.finfo(float).tiny, 1))
    selected = complete["gene_symbol"].isin(display["gene_symbol"])
    axb.scatter(px[~selected], py[~selected], s=8, color="#AEB4B9", alpha=0.6, linewidths=0, rasterized=True)
    axb.scatter(px[selected], py[selected], s=18,
                c=[vermilion if value > 0 else navy for value in px[selected]],
                edgecolor="white", linewidth=0.35, zorder=3)
    axb.axvline(0, color=charcoal, lw=0.65)
    axb.set_xlabel("Pooled Hedges' g")
    axb.set_ylabel(r"$-\log_{10}$ modified HK P")
    axb.set_title("All complete three-source genes", loc="left", pad=5)
    axb.grid(color=lightgrey, linewidth=0.45, zorder=0)
    axb.text(0.97, 0.96, "0 / 783\nBH FDR < 0.05", transform=axb.transAxes, ha="right", va="top",
             fontsize=6.8, fontweight="bold", color=charcoal,
             bbox=dict(boxstyle="round,pad=0.35", fc="white", ec=midgrey, lw=0.7))
    panel_label(axb, "b", x=-0.17)
    save_main(fig, "Figure_3")

    # Figure 4: heatmap with symmetric zero-centred color mapping and a separate
    # replication-call panel, avoiding reliance on color or asterisks alone.
    p = pd.read_csv(TABLES / "canonical_pathway_permutation_results.csv")
    p = p.loc[p["analysis_tier"].eq("primary glomerular meta-analysis")]
    rep = pd.read_csv(TABLES / "primary_glomerular_pathway_replication_summary.csv")
    pathway_order = rep["reactome_name"].tolist()
    pivot = p.pivot(index="reactome_name", columns="cohort", values="observed_mean_hedges_g").loc[pathway_order]
    pivot = pivot[["GSE96804", "GSE30528", "GSE104948_H7"]]
    fwer = p.pivot(index="reactome_name", columns="cohort", values="maxT_fwer_p").loc[pivot.index, pivot.columns]
    short_pathway = {
        "Complement cascade": "Complement cascade",
        "Coagulation pathway": "Coagulation pathway",
        "Cell surface interactions at the vascular wall": "Vascular-wall interactions",
        "Chemokine receptors bind chemokines": "Chemokine–receptor binding",
        "Extracellular matrix organization": "Extracellular-matrix organization",
        "Cellular response to hypoxia": "Cellular response to hypoxia",
        "Signaling by TGF-beta Receptor Complex": "TGF-β receptor signaling",
    }
    vmax = max(0.75, float(np.nanquantile(np.abs(pivot.to_numpy()), 0.99)))
    signed_cmap = LinearSegmentedColormap.from_list("cvd_signed", [navy, "#FFFFFF", vermilion])
    fig = plt.figure(figsize=(7.20, 4.55), constrained_layout=True)
    gs = fig.add_gridspec(1, 2, width_ratios=[1.72, 0.72], wspace=0.08)
    ax = fig.add_subplot(gs[0, 0])
    axr = fig.add_subplot(gs[0, 1], sharey=ax)
    im = ax.imshow(pivot.to_numpy(), cmap=signed_cmap, vmin=-vmax, vmax=vmax, aspect="auto")
    ax.set_xticks(range(3), ["GSE96804", "GSE30528", "GSE104948 H7"])
    ax.set_yticks(range(len(pivot)), [short_pathway[x] for x in pivot.index])
    ax.tick_params(length=0)
    for i in range(len(pivot)):
        for j in range(3):
            sig = fwer.iloc[i, j] < .05
            text_color = "white" if abs(pivot.iloc[i, j]) > vmax * 0.56 else charcoal
            ax.text(j, i, f"{pivot.iloc[i,j]:+.2f}\nP={fwer.iloc[i,j]:.3g}",
                    ha="center", va="center", fontsize=6.2, color=text_color,
                    fontweight="bold" if sig else "normal")
            if sig:
                ax.scatter(j + 0.36, i - 0.29, s=12, marker="*", color=text_color,
                           edgecolors="none", zorder=4)
    ax.set_title("Study-wise pathway evidence", loc="left", pad=6)
    for spine in ax.spines.values():
        spine.set_visible(False)
    cbar = fig.colorbar(im, ax=ax, orientation="horizontal", location="bottom", pad=0.09,
                        fraction=0.06, aspect=24)
    cbar.set_label("Mean Hedges' g (DKD − control)", fontsize=6.6)
    cbar.ax.tick_params(labelsize=6, length=2)
    panel_label(ax, "a", x=-0.42)

    rep = rep.set_index("reactome_name").loc[pivot.index]
    ry = np.arange(len(rep))
    axr.axvspan(1.5, 3.0, color="#E9F4EF", zorder=0)
    axr.hlines(ry, 0, rep["n_maxT_fwer_lt_0_05"], color="#A3A9AE", lw=1.0, zorder=1)
    axr.scatter(rep["n_maxT_fwer_lt_0_05"], ry,
                c=[green if bool(v) else midgrey for v in rep["primary_replication_call"]],
                s=28, edgecolor="white", linewidth=0.5, zorder=3)
    for i, row in enumerate(rep.itertuples()):
        axr.text(2.88, i, "Replicated" if row.primary_replication_call else "Not replicated",
                 ha="right", va="center", fontsize=5.8,
                 color=green if row.primary_replication_call else charcoal,
                 fontweight="bold" if row.primary_replication_call else "normal")
    axr.axvline(2, color=green, ls="--", lw=0.7)
    axr.set_xlim(-0.2, 3.05)
    axr.set_xticks([0, 1, 2, 3])
    axr.set_xlabel("Sources with maxT P < 0.05")
    axr.set_title("Replication rule", loc="left", pad=6)
    axr.tick_params(axis="y", left=False, labelleft=False)
    axr.grid(axis="x", color=lightgrey, linewidth=0.45)
    panel_label(axr, "b", x=-0.20)
    save_main(fig, "Figure_4")

    prov = pd.read_csv(TABLES / "canonical_pathway_provenance.csv")
    fig, ax = plt.subplots(figsize=(8.5, 5.2)); ax.barh(prov["reactome_name"], prov["human_gene_count_after_GOA_filter"], color="#4C8798")
    ax.set_xlabel("Canonical human genes"); ax.set_title("Supplementary Figure S1. Fixed Reactome pathway sizes", weight="bold"); fig.tight_layout()
    fig.savefig(SUPP / "Supplementary_Figure_S1.png", dpi=350, bbox_inches="tight"); plt.close(fig)
    cov = pd.read_csv(TABLES / "covariate_adjustment_pathway_sensitivity.csv")
    fig, ax = plt.subplots(figsize=(7.5,6)); ax.scatter(cov["mean_standardized_beta_unadjusted"], cov["mean_standardized_beta_adjusted"], c=["#D05A47" if x else "#24557A" for x in cov["direction_changed"]]); lim=max(abs(cov[["mean_standardized_beta_unadjusted","mean_standardized_beta_adjusted"]]).max().max(),.1); ax.plot([-lim,lim],[-lim,lim],"--",color="grey"); ax.set(xlabel="Unadjusted mean standardized beta",ylabel="Adjusted mean standardized beta",title="Supplementary Figure S2. Measured-covariate sensitivity"); fig.tight_layout(); fig.savefig(SUPP/"Supplementary_Figure_S2.png",dpi=350,bbox_inches="tight"); plt.close(fig)
    probe = pd.read_csv(TABLES / "probe_aggregation_pathway_sensitivity.csv")
    fig, ax=plt.subplots(figsize=(7.5,6)); ax.scatter(probe["mean_hedges_g_highest_mean_probe"],probe["mean_hedges_g_median_probe"],c="#6D5A8D"); lim=max(abs(probe[["mean_hedges_g_highest_mean_probe","mean_hedges_g_median_probe"]]).max().max(),.1); ax.plot([-lim,lim],[-lim,lim],"--",color="grey"); ax.set(xlabel="Highest-mean probe",ylabel="Median-probe aggregation",title="Supplementary Figure S3. Probe-aggregation sensitivity"); fig.tight_layout(); fig.savefig(SUPP/"Supplementary_Figure_S3.png",dpi=350,bbox_inches="tight"); plt.close(fig)


def build_source_tables() -> dict[str, Path]:
    mapping = {
        "S1": "systematic_dataset_screening_m19.csv", "S2": "cohort_characteristics.csv",
        "S3": "canonical_pathway_provenance.csv", "S4": "primary_glomerular_gene_meta.csv",
        "S5": "per_cohort_canonical_gene_effects.csv", "S6": "canonical_pathway_permutation_results.csv",
        "S7": "primary_glomerular_pathway_replication_summary.csv", "S8": "canonical_pathway_replication_summary.csv",
        "S9": "canonical_pathway_leave_one_gene_out.csv", "S10": "canonical_pathway_leave_one_sample_out.csv",
        "S11": "covariate_adjustment_gene_sensitivity.csv", "S12": "covariate_adjustment_pathway_sensitivity.csv",
        "S13": "sample_correlation_qc.csv", "S14": "probe_aggregation_sensitivity.csv",
        "S15": "probe_aggregation_pathway_sensitivity.csv", "S16": "small_sample_glomerular_gene_meta_sensitivity.csv",
        "S17": "descriptive_tubulointerstitial_gene_meta.csv", "S18": "descriptive_whole_kidney_gene_meta.csv",
    }
    paths = {}
    for sid, name in mapping.items():
        source = TABLES / name; target = SUPP / f"Supplementary_Table_{sid}.csv"; shutil.copy2(source,target); paths[sid]=target
    search_rows=[]
    for path in sorted((DOCS/"systematic_search").glob("*")):
        if path.is_file(): search_rows.append({"file":path.name,"size_bytes":path.stat().st_size,"sha256":sha256(path),"role":"systematic search evidence"})
    for path in sorted((ROOT/"data_raw"/"reference_gene_sets_m19").glob("*")):
        if path.is_file(): search_rows.append({"file":path.name,"size_bytes":path.stat().st_size,"sha256":sha256(path),"role":"canonical gene-set source"})
    target=SUPP/"Supplementary_Table_S19.csv"; pd.DataFrame(search_rows).to_csv(target,index=False); paths["S19"]=target
    return paths


DESCRIPTIONS = {
 "S1":"All 322 unique dataset records with source, screening stage, decision, and specific reason.",
 "S2":"Included GEO Series, source-study identity, compartment, platform, group sizes, controls, preprocessing, covariates, and analysis tier.",
 "S3":"Fixed Reactome pathway identifiers, memberships, source hashes, GOA filtering, and retrieval date.",
 "S4":"Primary three-source glomerular REML meta-analysis for the 783-gene canonical union; nonestimable members retained in multiplicity bookkeeping.",
 "S5":"Complete per-cohort Hedges' g, variance, Welch P value, and within-cohort canonical-family FDR.",
 "S6":"Two-sided joint-label permutation results and seven-pathway maxT FWER for every cohort.",
 "S7":"Primary glomerular pathway replication calls under the fixed two-of-three concordant maxT rule.",
 "S8":"All-tier compartment-specific pathway evidence summary; not used for primary replication calls.",
 "S9":"Leave-one-gene-out pathway mean-effect sensitivity.", "S10":"Leave-one-sample-out pathway mean-effect sensitivity; cohorts with fewer than three per group marked non-evaluable.",
 "S11":"Gene-level measured-covariate OLS sensitivity for datasets with usable sex and/or age metadata.", "S12":"Pathway summaries of measured-covariate sensitivity.",
 "S13":"Within-cohort sample-correlation screening; flags are not automatic exclusions.", "S14":"Gene-level highest-mean versus median-probe sensitivity for GSE30528/GSE30529.",
 "S15":"Pathway-level probe-aggregation sensitivity and canonical-union gene-effect correlations.", "S16":"Glomerular gene meta-analysis including donor-averaged GSE1009 as a small-sample sensitivity.",
 "S17":"Descriptive two-source tubulointerstitial gene synthesis.", "S18":"Descriptive whole/cortical-kidney gene synthesis including small studies.",
 "S19":"Search-evidence and canonical-reference file hashes."
}


def build_workbook(paths: dict[str, Path]) -> Path:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    wb=Workbook(); ws=wb.active; ws.title="README"; ws.append(["Table","Definition"])
    for sid in paths: ws.append([f"Table_{sid}",DESCRIPTIONS[sid]])
    dictionary=[]
    for sid,path in paths.items():
        frame=pd.read_csv(path,nrows=10)
        sheet=wb.create_sheet(f"Table_{sid}")
        full=pd.read_csv(path)
        sheet.append(list(full.columns))
        for row in full.itertuples(index=False,name=None): sheet.append(list(row))
        for col in full.columns:
            dictionary.append([sid,col,str(full[col].dtype),f"Field `{col}` in {DESCRIPTIONS[sid]}"])
    d=wb.create_sheet("Data_dictionary"); d.append(["table","field","type","definition"])
    for row in dictionary:d.append(row)
    for sheet in wb.worksheets:
        sheet.freeze_panes="A2"; sheet.auto_filter.ref=sheet.dimensions
        for cell in sheet[1]: cell.font=Font(bold=True,color="FFFFFF"); cell.fill=PatternFill("solid",fgColor="24557A"); cell.alignment=Alignment(wrap_text=True)
        for col in sheet.columns:
            letter=col[0].column_letter; sheet.column_dimensions[letter].width=min(max(12,max(len(str(c.value or "")) for c in col[:100])+2),45)
    out=SUPP/"Supplementary_Tables_S1-S19.xlsx"; wb.save(out); return out


def manuscript_text() -> str:
    cohorts=pd.read_csv(TABLES/"cohort_characteristics.csv").copy(); cohorts["case_control"]=cohorts["n_case"].astype(str)+" / "+cohorts["n_control"].astype(str)
    rep=pd.read_csv(TABLES/"primary_glomerular_pathway_replication_summary.csv").copy(); rep["call"]=rep["primary_replication_call"].map({True:"Replicated",False:"Not replicated"})
    table1=markdown_table(cohorts,["cohort","source_study","compartment","case_control","analysis_tier"],["Series/analysis","Source study","Compartment","DKD/control","Role"])
    table2=markdown_table(rep,["reactome_name","n_positive","n_maxT_fwer_lt_0_05","call"],["Reactome pathway","Positive sources","Sources with maxT P<0.05","Primary call"])
    return f"""# Compartment-Stratified Systematic Reanalysis of Complement, Coagulation, and Matrix Transcriptional Programs in Diabetic Kidney Disease

Lizhen Deng¹*  
¹College of Life Science and Technology, Huazhong University of Science and Technology, Wuhan, Hubei, China  
*Correspondence: Lizhen Deng; 3070116993@qq.com; ORCID 0009-0003-2428-8176

## Abstract

Public diabetic kidney disease (DKD) transcriptomes differ in tissue compartment, platform, control source, and clinical annotation, complicating claims of molecular replication. We systematically searched GEO, ArrayExpress, and PubMed through 21 August 2026 and screened 322 unique dataset records. Eleven GEO Series representing nine source studies met eligibility criteria. The primary analysis was restricted to three independent glomerular sources and a fixed seven-pathway Reactome family comprising 783 human genes; other compartments were analyzed separately. Random-effects synthesis used restricted maximum likelihood and modified Hartung–Knapp inference, with Benjamini–Hochberg correction across all 783 genes. No gene met FDR<0.05. Joint two-sided sample-label permutations with maxT family-wise correction identified reproducible positive net effects for complement cascade, chemokine-receptor binding, and extracellular-matrix organization in at least two of three glomerular sources. Coagulation, hypoxia, vascular-wall interaction, and TGF-β signaling did not meet the replication rule. Sex/age adjustment, leave-one-sample/gene analyses, and probe-aggregation sensitivities qualified but did not overturn the principal conclusions. These results support compartment-bounded transcriptional convergence of selected inflammatory and matrix programs, not causal pathway activation or a replicated coagulation signature.

## Keywords

diabetic kidney disease; glomerulus; complement; extracellular matrix; transcriptomics; meta-analysis

## Introduction

Diabetic kidney disease (DKD) combines glomerular, tubular, vascular, inflammatory, and fibrotic injury [1]. Human kidney transcriptomic studies have reported recurrent immune and matrix-associated expression changes, but the available cohorts differ in renal compartment, disease stage, control tissue, assay, and preprocessing [2–5]. Pooling these settings without a defined biological estimand can produce a precise number whose biological meaning is unclear.

Earlier versions of this project used manually assembled gene groups and reused datasets across discovery and synthesis. The present analysis replaces those groups with a fixed, externally versioned Reactome family [14], conducts a reproducible repository and literature search following PRISMA reporting principles [15], counts source studies rather than accessions when donor or parent-study overlap is present, and prohibits cross-compartment pooling. We asked two bounded questions: whether individual genes in the canonical family show multiplicity-controlled evidence across three independent glomerular sources, and whether canonical pathway mean effects recur under within-study joint-label tests with family-wise error control.

## Results

### Systematic search and eligible evidence

The searches identified 263 GEO records, 54 ArrayExpress records, and 57 GSE accessions through 1,088 PubMed records used solely for accession discovery. After 52 database duplicates were removed, 322 unique dataset records were screened; 269 were excluded at title/summary screening and 53 underwent full accession and sample-metadata review. Forty-two were excluded, leaving 11 Series representing nine source studies (Fig. 1; Supplementary Table S1). GSE30528/GSE30529 are compartments of GSE30122; GSE104948/GSE104954 are compartments of ERCB H7. GSE99339/GSE99325 were excluded because their participant/source labels overlap GSE104948/GSE104954.

### Compartment-specific design

The primary glomerular synthesis contained GSE96804 (41 DKD/20 controls), GSE30528 (9/13), and the H7 stratum of GSE104948 (7/18). The H7 restriction avoided mixing batches. Tubulointerstitial, whole/cortical-kidney, interstitium-only, and small-sample cohorts were evaluated separately (Fig. 2; Table 1). No effect was pooled across compartments.

### Primary gene-level synthesis

The seven canonical pathways contained 783 unique GOA-filtered human symbols. Complete estimates from all three primary glomerular sources were available for 582 genes. REML random-effects models with modified Hartung–Knapp inference yielded no gene meeting BH FDR<0.05 across the 783-member family (0/783; Fig. 3; Supplementary Table S4). This null multiplicity-controlled result limits individual-gene claims; it does not establish absence of all expression differences.

### Canonical pathway replication

Each cohort used the mean unaligned Hedges' g over mapped genes. Case/control labels were reassigned jointly across samples, preserving the observed gene-correlation structure. Exact enumeration was used when there were no more than 10,000 allocations; otherwise 10,000 Monte Carlo permutations and plus-one P values were used. maxT controlled family-wise error across the fixed seven pathways. Complement cascade, chemokine receptors bind chemokines, and extracellular matrix organization were positive in all three primary sources and had maxT P<0.05 in two sources, satisfying the prespecified replication rule (Fig. 4; Table 2). Coagulation was mixed in direction and significant in no primary source after maxT correction. Vascular-wall interactions were positive in all three but significant in only one; hypoxia and TGF-β were directionally heterogeneous.

### Sensitivity analyses

Sex adjustment in GSE96804 did not change any pathway direction. Age-plus-sex adjustment in GSE166239 changed one of seven directions; sex adjustment in GSE163603 changed none, although all DKD samples there were male and residual confounding remains. Leave-one-sample-out analysis produced no direction changes in the three primary glomerular cohorts; eight changes occurred only in contextual cohorts. Highest-mean versus median-probe aggregation correlated strongly across canonical-union gene effects (Pearson 0.913 in GSE30528 and 0.901 in GSE30529); only the nonreplicated coagulation pathway changed direction in GSE30528. Correlation flags were retained as screening flags, not automatic sample exclusions, because all could reflect disease-associated global structure rather than technical failure.

### Table 1. Eligible Series and analysis roles

{table1}

### Table 2. Primary glomerular pathway replication

{table2}

## Discussion

The principal result is a deliberate separation of gene-level and pathway-level evidence. No member of the 783-gene canonical union survived the small-k, multiplicity-controlled glomerular meta-analysis. Nevertheless, three broad programs showed study-wise, family-wise-controlled positive mean effects in at least two independent glomerular sources. Aggregation can stabilize weak, distributed effects, but it does not identify which genes are causal, whether proteins or biochemical pathway activity change, or which cell type produces the signal.

Complement and matrix remodeling are biologically compatible with DKD inflammation and scarring [1–5]. The chemokine-receptor result similarly supports a distributed inflammatory transcript pattern. By contrast, the coagulation pathway did not replicate: its net direction differed across primary studies and none passed maxT correction. The title retains coagulation because it was a fixed target of the reanalysis, but the conclusion is explicitly negative for a replicated coagulation signature.

Compartment discipline materially changed the interpretation. Glomeruli, tubulointerstitium, whole biopsy/cortex, and isolated interstitium have different cellular compositions and control sources. GSE30528 and GSE30529 cannot be counted as independent evidence, and GSE104948/GSE104954 likewise represent paired compartments from one source collection. Separate contextual results are therefore descriptive rather than components of a universal DKD effect.

The study remains limited by archived observational data. Sample-level age, sex, renal function, medication, batch, and ancestry were incompletely recorded. Label permutation is exact or Monte Carlo-valid only under exchangeability of the archived labels and cannot remove confounding. Control tissues include living-donor and nephrectomy sources. Platforms and upstream preprocessing differ, and the analysis begins from processed matrices or normalized GEO values rather than uniform raw-file reprocessing. Only three independent sources inform the primary gene synthesis, making heterogeneity and interval estimates imprecise [9,10]. The pathway family was fixed for this revision but not prospectively preregistered. Systematic-search screening was performed by one reviewer. Finally, overlapping pathway membership means that a positive mean effect is not evidence of independent mechanisms.

## Methods

### Search strategy and eligibility

GEO DataSets was searched through NCBI E-utilities with two field-qualified human GSE queries covering diabetic nephropathy/DKD and broader diabetes-plus-kidney terminology. The ArrayExpress collection in BioStudies was searched for “diabetic nephropathy” and “diabetic kidney disease”; its synonym expansion returned the same 54 records. PubMed was searched with two title/abstract expressions combining DKD terms, transcriptomic terms, kidney terms, and human/patient terms. All query strings, translations, timestamps, returned identifiers, summaries, and screening decisions are archived. PubMed articles were used only to discover repository accessions, so article counts were not mixed with dataset-record counts. Eligible studies contained human bulk kidney-tissue expression data with separable DKD and non-DKD groups. Cell culture, animal-only, blood/urine, miRNA-only, single-cell/nucleus, spatial, noncase-control, derived duplicate, and unresolvable-overlap records were excluded. Small studies were retained as sensitivities rather than excluded by an arbitrary size threshold [15].

### Canonical pathway family

Seven Reactome pathways were fixed before the M19 calculations: Complement cascade (R-HSA-166658), Coagulation pathway (R-HSA-9769740), Cell surface interactions at the vascular wall (R-HSA-202733), Chemokine receptors bind chemokines (R-HSA-380108), Extracellular matrix organization (R-HSA-1474244), Cellular response to hypoxia (R-HSA-1234174), and Signaling by TGF-beta Receptor Complex (R-HSA-170834). The 21 June 2026 Reactome GMT was filtered to symbols in the 28 May 2026 GOA human GAF, yielding 783 unique genes. Source URLs, retrieval dates, SHA-256 hashes, and memberships are archived [14]. Historical manually curated sets have no confirmatory role.

### Expression processing and study effects

Expression matrices were analyzed within cohorts. Microarray probes were mapped to uppercase gene symbols; the primary label-independent aggregation selected the probe with highest all-sample mean, with median-probe aggregation as a sensitivity. Author count tables were transformed to log2(CPM+0.5). GSE1009 apparent a/b technical replicates were averaged at the donor-label level. Hedges' g represented DKD minus control standardized mean difference [8], and two-sided Welch P values were retained for within-study description.

### Gene meta-analysis

The primary estimand was the mean glomerular DKD-control standardized difference across GSE96804, GSE30528, and GSE104948 H7. Each gene used REML between-study variance and modified Hartung–Knapp t inference with k−1 degrees of freedom [9,10]. Nonestimable genes remained in the 783-member family with P=1 for BH correction [11]. No cross-compartment estimate was calculated.

### Pathway permutation and replication rule

For each pathway and cohort, the statistic was the arithmetic mean of observed gene-level Hedges' g without aligning signs to an outcome-derived direction. All gene effects were recomputed jointly after reallocating the fixed number of case labels. When the number of allocations was ≤10,000 every allocation was enumerated; otherwise 10,000 allocations were sampled with seed values derived from 20260821. Two-sided plus-one P values and maxT FWER were reported over seven pathways. A pathway was called replicated only if at least two of the three independent primary glomerular sources had maxT FWER P<0.05 with the same net direction.

### Robustness and quality control

Measured-covariate OLS sensitivities used disease plus sex for GSE96804 and GSE163603, and disease plus standardized age plus sex for GSE166239. These models were not generalized to datasets lacking comparable metadata. Leave-one-gene and leave-one-sample analyses assessed pathway sign stability. Sample QC used median Pearson correlation over the 1,000 most variable genes and flagged robust z<−3. Probe aggregation was compared in GSE30528/GSE30529. Flags were not exclusion rules without independent technical evidence.

### Reproducibility and AI assistance

Scripts, search logs, fixed pathway files, source tables, checksums, and an environment specification accompany the submission. OpenAI Codex was used under author supervision for code generation, statistical comparison, consistency checking, figures, tables, and language drafting. No generative image model was used; figures were generated from archived numerical tables. The author is responsible for verifying all code, results, references, and prose. AI was not used to generate or alter primary data and is not an author.

## Data availability

All input expression data are publicly available in GEO under GSE1009, GSE30528, GSE30529, GSE96804, GSE104948, GSE104954, GSE111154, GSE142025, GSE163603, GSE166239, and GSE199838. Accession-level eligibility, sample definitions, processed effects, and hashes are supplied in Supplementary Tables S1–S19. No new primary data were generated.

## Code availability

The exact M19 code, search logs, tables, manuscript, and submission assets are publicly archived at https://github.com/denglizhen-113/dkd-focused-universe-reanalysis/tree/v1.1.2 and are also supplied in Source_Code_M19.zip and Source_Data_M19.zip. The Git tag provides a versioned public snapshot; no DOI has been assigned.

## References

1. Alicic, R. Z., Rooney, M. T. & Tuttle, K. R. Diabetic kidney disease: challenges, progress, and possibilities. *Clin. J. Am. Soc. Nephrol.* **12**, 2032–2045 (2017). https://doi.org/10.2215/CJN.11491116
2. Woroniecka, K. I. et al. Transcriptome analysis of human diabetic kidney disease. *Diabetes* **60**, 2354–2369 (2011). https://doi.org/10.2337/db10-1181
3. Pan, Y. et al. Dissection of glomerular transcriptional profile in patients with diabetic nephropathy: SRGAP2a protects podocyte structure and function. *Diabetes* **67**, 717–730 (2018). https://doi.org/10.2337/db17-0755
4. Fan, Y. et al. Comparison of kidney transcriptomic profiles of early and advanced diabetic nephropathy reveals potential new mechanisms for disease progression. *Diabetes* **68**, 2301–2314 (2019). https://doi.org/10.2337/db19-0204
5. Nordbø, O. P. et al. Transcriptomic analysis reveals partial epithelial-mesenchymal transition and inflammation as common pathogenic mechanisms in hypertensive nephrosclerosis and type 2 diabetic nephropathy. *Physiol. Rep.* **11**, e15825 (2023). https://doi.org/10.14814/phy2.15825
6. Squair, J. W. et al. Confronting false discoveries in single-cell differential expression. *Nat. Commun.* **12**, 5692 (2021). https://doi.org/10.1038/s41467-021-25960-2
7. Lähnemann, D. et al. Eleven grand challenges in single-cell data science. *Genome Biol.* **21**, 31 (2020). https://doi.org/10.1186/s13059-020-1926-6
8. Hedges, L. V. Distribution theory for Glass's estimator of effect size and related estimators. *J. Educ. Stat.* **6**, 107–128 (1981). https://doi.org/10.2307/1164588
9. IntHout, J., Ioannidis, J. P. A. & Borm, G. F. The Hartung-Knapp-Sidik-Jonkman method for random effects meta-analysis. *BMC Med. Res. Methodol.* **14**, 25 (2014). https://doi.org/10.1186/1471-2288-14-25
10. Röver, C., Knapp, G. & Friede, T. Hartung-Knapp-Sidik-Jonkman approach and its modification for random-effects meta-analysis with few studies. *BMC Med. Res. Methodol.* **15**, 99 (2015). https://doi.org/10.1186/s12874-015-0091-1
11. Benjamini, Y. & Hochberg, Y. Controlling the false discovery rate. *J. R. Stat. Soc. B* **57**, 289–300 (1995). https://doi.org/10.1111/j.2517-6161.1995.tb02031.x
12. Langfelder, P. & Horvath, S. WGCNA: an R package for weighted correlation network analysis. *BMC Bioinformatics* **9**, 559 (2008). https://doi.org/10.1186/1471-2105-9-559
13. Ritchie, M. E. et al. limma powers differential expression analyses. *Nucleic Acids Res.* **43**, e47 (2015). https://doi.org/10.1093/nar/gkv007
14. Gillespie, M. et al. The Reactome pathway knowledgebase 2022. *Nucleic Acids Res.* **50**, D687–D692 (2022). https://doi.org/10.1093/nar/gkab1028
15. Page, M. J. et al. The PRISMA 2020 statement. *BMJ* **372**, n71 (2021). https://doi.org/10.1136/bmj.n71

## Acknowledgements

The author acknowledges the investigators who generated and publicly shared the reanalyzed datasets.

## Funding

The author received no specific funding for this work.

## Ethics statement

This secondary analysis used only publicly available, de-identified transcriptomic data and involved no new recruitment, intervention, specimen collection, participant contact, or access to identifiable private information. No new participant consent was sought. Ethics and consent procedures for original collection remain the responsibility of the source studies and are described in their repository records and publications. No animal data were analyzed.

## Author contributions

Lizhen Deng: conceptualization, methodology, software, formal analysis, investigation, data curation, visualization, writing—original draft, writing—review and editing, and project administration.

## Competing interests

The author declares no competing interests.

## Figure legends

**Figure 1. Systematic dataset identification and eligibility.** Dataset records, rather than PubMed articles, are the screening unit. The right-hand branches show records removed before screening and exclusions at the title/summary and full-record stages. Full record-level decisions are provided in Supplementary Table S1.

**Figure 2. Compartment-specific study architecture.** **a,** Mirrored bars show control and DKD group sizes; color denotes renal compartment. **b,** Each row is an independent source study and each symbol is an analyzed GEO Series. Symbol shape denotes the primary, contextual, or sensitivity role; symbol area scales with total sample size. GSE30528/GSE30529 and GSE104948 H7/GSE104954 H7 occupy paired compartment cells on single source-study rows and were not treated as independent source effects.

**Figure 3. Primary glomerular gene synthesis.** **a,** Pooled Hedges' g and 95% modified Hartung–Knapp (HK) confidence intervals for the 12 genes with the lowest unadjusted modified-HK P values among complete three-source estimates; selection is descriptive and is not evidence of multiplicity-controlled significance. **b,** Pooled effects and unadjusted modified-HK P values for all 582 genes with complete estimates across the three primary glomerular sources. Benjamini–Hochberg correction was applied across the full 783-gene canonical family, including nonestimable members assigned P=1; no gene met FDR<0.05.

**Figure 4. Canonical pathway evidence in three primary glomerular sources.** **a,** Cells show the unaligned mean Hedges' g (DKD minus control) and two-sided maxT family-wise-error-rate P value for each pathway and source. The color scale is symmetric about zero; asterisks and bold text indicate maxT P<0.05. **b,** Number of primary sources meeting maxT P<0.05. The dashed line marks the prespecified two-source threshold; a pathway was called replicated only when at least two sources were significant in the same net direction.
"""


def supplement_text() -> str:
    rows="\n".join(f"- **Supplementary Table {sid}.** {DESCRIPTIONS[sid]}" for sid in DESCRIPTIONS)
    return f"""# Supplementary Information

## Study title

Compartment-Stratified Systematic Reanalysis of Complement, Coagulation, and Matrix Transcriptional Programs in Diabetic Kidney Disease

## Supplementary methods and audit boundaries

This supplement accompanies the complete machine-readable source tables. The primary evidence is restricted to three independent glomerular source studies. Contextual compartments, small cohorts, correlation flags, covariate models, and alternative probe aggregation are sensitivities. No cross-compartment pooled effect is reported. Exact search queries and translations are included in the source-data archive.

## Supplementary figures

### Supplementary Figure S1. Fixed Reactome pathway sizes

![Supplementary Figure S1](Supplementary_Figure_S1.png)

Gene counts after filtering the current Reactome GMT to current GOA human symbols.

### Supplementary Figure S2. Measured-covariate sensitivity

![Supplementary Figure S2](Supplementary_Figure_S2.png)

Unadjusted versus adjusted pathway mean standardized disease coefficients. Red denotes a direction change.

### Supplementary Figure S3. Probe-aggregation sensitivity

![Supplementary Figure S3](Supplementary_Figure_S3.png)

Pathway mean Hedges' g using the highest-all-sample-mean probe versus the median across mapped probes.

## Supplementary tables

{rows}
"""


def checklist_text() -> str:
    items=[
      ("1","Title identifies a systematic reanalysis","Title page"),("2","Structured summary","Abstract"),("3","Rationale","Introduction"),("4","Objectives","Introduction, final paragraph"),("5","Eligibility criteria","Methods: Search strategy and eligibility"),("6","Information sources and last search date","Methods; Supplementary Table S19"),("7","Full search strategies","Source_Data_M19.zip search logs"),("8","Selection process","Methods; Figure 1; Supplementary Table S1"),("9","Data collection process","Methods: Expression processing"),("10a","Outcomes and variables","Methods"),("11","Risk of bias/confounding assessment","Robustness/QC and Limitations"),("12","Effect measures","Methods: Hedges' g"),("13a","Synthesis eligibility","Methods"),("13d","Sensitivity analyses","Results and Supplementary Tables S9-S15"),("14","Reporting bias","Not formally estimable with three primary studies; stated in Limitations"),("15","Certainty assessment","No formal GRADE; bounded inferential language in Discussion"),("16a","Study selection","Figure 1"),("17","Study characteristics","Table 1; Supplementary Table S2"),("19","Individual study results","Figure 4; Supplementary Tables S5-S6"),("20","Synthesis results","Results; Figures 3-4"),("22","Certainty","Discussion and Limitations"),("23","Discussion","Discussion"),("24","Registration/protocol","Not registered; stated explicitly"),("25","Support","Funding"),("26","Competing interests","Competing interests"),("27","Availability","Data and Code availability")]
    body="\n".join(f"| {a} | {b} | {c} |" for a,b,c in items)
    return "# PRISMA 2020 Checklist\n\n| Item | Requirement | Location/reporting |\n| --- | --- | --- |\n"+body+"\n"


def add_page_number(paragraph) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER; run=paragraph.add_run()
    for kind,text in (("begin",None),(None," PAGE "),("separate",None),("end",None)):
        if text is not None:
            el=OxmlElement("w:instrText"); el.set(qn("xml:space"),"preserve"); el.text=text
        else:
            el=OxmlElement("w:fldChar"); el.set(qn("w:fldCharType"),kind)
        run._r.append(el)


def convert(md: Path, docx: Path, pdf: Path, kind="manuscript") -> None:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
    subprocess.run([str(PANDOC),str(md),"--from=gfm","--to=docx","--standalone",f"--resource-path={SUPP};{PACKAGE};{ROOT}","--output",str(docx)],check=True,cwd=ROOT)
    doc=Document(docx); normal=doc.styles["Normal"]; normal.font.name="Times New Roman"; normal.font.size=Pt(11); normal._element.rPr.rFonts.set(qn("w:eastAsia"),"Times New Roman")
    for section in doc.sections:
        section.top_margin=section.bottom_margin=Inches(.85); section.left_margin=section.right_margin=Inches(.9); add_page_number(section.footer.paragraphs[0])
        if kind=="manuscript":
            ln=OxmlElement("w:lnNumType"); ln.set(qn("w:countBy"),"1"); ln.set(qn("w:restart"),"continuous"); section._sectPr.append(ln)
    doc.save(docx)
    subprocess.run(["powershell","-NoProfile","-ExecutionPolicy","Bypass","-File",str(PDF_EXPORTER),"-DocxPath",str(docx),"-PdfPath",str(pdf)],check=True,cwd=ROOT)


def main() -> None:
    for path in (DOCS,MANUSCRIPT,PACKAGE,SUPP,MAIN_FIG,READY): path.mkdir(parents=True,exist_ok=True)
    if "--figures-only" in sys.argv:
        build_figures()
        print("figures_built=4+3")
        return
    if "--skip-figures" not in sys.argv:
        build_figures()
    paths=build_source_tables(); workbook=build_workbook(paths)
    manuscript=MANUSCRIPT/"scientific_reports_m19_manuscript.md"; manuscript.write_text(manuscript_text(),encoding="utf-8")
    supp=SUPP/"supplementary_information.md"; supp.write_text(supplement_text(),encoding="utf-8")
    cover=PACKAGE/"cover_letter_scientific_reports.md"; cover.write_text("""# Cover Letter\n\n21 August 2026\n\nEditors, *Scientific Reports*\n\nDear Editors,\n\nPlease consider “Compartment-Stratified Systematic Reanalysis of Complement, Coagulation, and Matrix Transcriptional Programs in Diabetic Kidney Disease.” This systematic reanalysis screens 322 unique dataset records, prevents double counting of compartments and overlapping source cohorts, uses a fixed Reactome family, and separates gene-level random-effects inference from study-wise maxT pathway evidence. No individual gene met FDR<0.05. Complement, chemokine-receptor binding, and extracellular-matrix organization—but not coagulation—met the explicit primary glomerular replication rule. The manuscript makes no causal, protein-activity, or universal cross-compartment claim.\n\nThe work is original and not under consideration elsewhere. The author declares no competing interests and no specific funding. Only public de-identified data were analyzed. Code, tables, search logs, and checksums accompany the submission and are publicly versioned at https://github.com/denglizhen-113/dkd-focused-universe-reanalysis/tree/v1.1.2. No DOI has been assigned.\n\nSincerely,\n\nLizhen Deng\nCollege of Life Science and Technology, Huazhong University of Science and Technology\nWuhan, Hubei, China\n3070116993@qq.com\nORCID 0009-0003-2428-8176\n""",encoding="utf-8")
    checklist=PACKAGE/"PRISMA_2020_checklist.md"; checklist.write_text(checklist_text(),encoding="utf-8")
    submission_check=PACKAGE/"SCIENTIFIC_REPORTS_SUBMISSION_CHECKLIST.md"; submission_check.write_text("""# Scientific Reports submission check\n\n- Title: 14 words; abstract: ≤200 words; keywords: 6.\n- Author, affiliation, correspondence, ORCID: populated from the prior author-approved package; recheck portal spelling.\n- Main manuscript: line-numbered DOCX; four figures uploaded separately.\n- Statistics: exact P values in source tables; two-sided tests; BH and maxT families stated.\n- Data/code: exact M19 archives attached and public tag v1.1.2 cited; no DOI has been assigned.\n- Ethics: public de-identified secondary analysis; original-study ethics remain with source studies.\n- Funding/competing interests/author contributions/AI assistance: included.\n- Human-verification fields still required at portal: institutional correspondence preference and suggested/excluded reviewers if requested.\n""",encoding="utf-8")
    approval=DOCS/"M19_AUTHOR_APPROVAL_RECORD.md"; approval.write_text("# M19 revision authorization\n\nOn 21 August 2026 the author instructed: “请你自行完善以上所有问题，我批准所有权限”. This records authorization to revise, calculate, assemble, and designate the completed M19 package. It does not authorize invention of ethics identifiers, reviewer identities, institutional email addresses, or repository DOI values.\n",encoding="utf-8")

    convert(manuscript,PACKAGE/"manuscript.docx",PACKAGE/"manuscript.pdf","manuscript")
    convert(supp,SUPP/"supplementary_information.docx",SUPP/"supplementary_information.pdf","supplement")
    convert(cover,PACKAGE/"cover_letter.docx",PACKAGE/"cover_letter.pdf","cover")
    convert(checklist,PACKAGE/"PRISMA_2020_checklist.docx",PACKAGE/"PRISMA_2020_checklist.pdf","checklist")

    with zipfile.ZipFile(SUPP/"Source_Data_M19.zip","w",zipfile.ZIP_DEFLATED) as z:
        for path in paths.values(): z.write(path,path.name)
        for path in sorted((DOCS/"systematic_search").glob("*")): z.write(path,f"systematic_search/{path.name}")
        z.write(TABLES/"prisma_counts_m19.json","prisma_counts_m19.json"); z.write(TABLES/"canonical_pathways.gmt","canonical_pathways.gmt")
    with zipfile.ZipFile(SUPP/"Source_Code_M19.zip","w",zipfile.ZIP_DEFLATED) as z:
        for path in sorted((ROOT/"scripts"/"stage21_m19_scientific_reports_revision").glob("*.py")): z.write(path,f"scripts/stage21_m19_scientific_reports_revision/{path.name}")
        z.write(ROOT/"environment_m18.yml","environment_m18.yml")
    ready_sources={"manuscript.docx":PACKAGE/"manuscript.docx","cover_letter.pdf":PACKAGE/"cover_letter.pdf","PRISMA_2020_checklist.pdf":PACKAGE/"PRISMA_2020_checklist.pdf","supplementary_information.pdf":SUPP/"supplementary_information.pdf","Supplementary_Tables_S1-S19.xlsx":workbook,"Source_Data_M19.zip":SUPP/"Source_Data_M19.zip","Source_Code_M19.zip":SUPP/"Source_Code_M19.zip"}
    for i in range(1,5): ready_sources[f"Figure_{i}.pdf"]=MAIN_FIG/f"Figure_{i}.pdf"
    for path in READY.iterdir():
        if path.is_file(): path.unlink()
    for name,source in ready_sources.items(): shutil.copy2(source,READY/name)
    manifest=pd.DataFrame([{"upload_file":n,"size_bytes":p.stat().st_size,"sha256":sha256(p)} for n,p in ready_sources.items()]); manifest.to_csv(PACKAGE/"final_upload_manifest.csv",index=False)
    print(json.dumps({"ready_files":len(ready_sources),"ready":str(READY)},ensure_ascii=False))


if __name__ == "__main__": main()
